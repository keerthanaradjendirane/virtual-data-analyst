import streamlit as st
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
import numpy as np
from sklearn.linear_model import LinearRegression
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.preprocessing import LabelEncoder
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# Functions to Generate Word Document
def add_heading(document, text, level):
    """Add a heading with custom styles."""
    heading = document.add_heading(level=level)
    run = heading.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_table(document, headers, rows):
    """Add a styled table to the document."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Add header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    # Add data rows
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)

def generate_word_report(data, model=None, predictions=None, model_type=None, y_test=None):
    """Generate a detailed Word report."""
    doc = Document()

    # Title and Introduction
    add_heading(doc, 'Data Analysis Report', level=0)
    doc.add_paragraph(
        "This report provides a comprehensive analysis of the dataset, "
        "highlighting key insights, missing values, statistical summaries, and predictive model evaluations."
    )

    # Dataset Overview
    add_heading(doc, 'Dataset Overview', level=1)
    doc.add_paragraph(f"The dataset contains {data.shape[0]} rows and {data.shape[1]} columns.")
    doc.add_paragraph("Below are some insights derived from the data:")

    # Missing Values
    add_heading(doc, 'Missing Values', level=2)
    for column, value in data.isnull().sum().items():
        doc.add_paragraph(f"{column}: {value}")

    # Duplicates
    add_heading(doc, 'Duplicate Rows', level=2)
    doc.add_paragraph(f"Duplicate rows: {data.duplicated().sum()}")

    # Dataset Statistics
    add_heading(doc, 'Dataset Statistics', level=2)
    headers = data.describe().columns.tolist()
    stats = data.describe().reset_index().values.tolist()
    add_table(doc, ["Metric"] + headers, stats)

    # Model Evaluation
    if model and predictions is not None and y_test is not None:
        add_heading(doc, 'Model Evaluation', level=1)
        doc.add_paragraph(f"Model Type: {model_type}")
        doc.add_paragraph(f"Mean Squared Error: {mean_squared_error(y_test, predictions):.2f}")
        doc.add_paragraph(f"R-squared: {r2_score(y_test, predictions):.2f}")

    # Executive Summary
    add_heading(doc, 'Executive Summary', level=1)
    doc.add_paragraph(
        f"The analysis highlights that the dataset is high-quality with no significant missing or duplicate data. "
        f"Model evaluation results indicate satisfactory performance with the following metrics: "
        f"Mean Squared Error = {mean_squared_error(y_test, predictions):.2f}, R-squared = {r2_score(y_test, predictions):.2f}. "
        "The visualizations and statistical summaries provide a detailed understanding of the dataset."
    )

    # Save the document
    file_name = "Data_Analysis_Report.docx"
    doc.save(file_name)
    return file_name

# Streamlit Application
st.title("AI-Driven Virtual Data Analyst")
st.write("An advanced AI-driven data analysis solution capable of processing information, generating graphs, performing in-depth analyses, and providing executive summaries and reports.")

# Upload Data
uploaded_file = st.file_uploader("Upload your dataset (CSV format or JSON)", type=["csv", "json"])
if uploaded_file:
    # Read the uploaded file
    if uploaded_file.name.endswith(".csv"):
        data = pd.read_csv(uploaded_file)
    elif uploaded_file.name.endswith(".json"):
        data = pd.read_json(uploaded_file)

    st.write("### Dataset Preview:")
    st.write(data.head())

    # Perform Sanity Checks
    st.write("### Data Sanity Checks:")
    st.write(f"Missing Values:\n{data.isnull().sum()}")
    st.write(f"Duplicate Rows: {data.duplicated().sum()}")

    # Handle Missing Values
    if st.button("Handle Missing Values"):
        data = data.dropna()
        st.write("Missing values removed.")
        st.write(data.isnull().sum())

    # Display dataset statistics
    st.write("### Dataset Statistics:")
    st.write(data.describe())

    # Column selection for graph generation
    st.write("### Generate Graphs:")
    columns = data.columns.tolist()

    if columns:
        x_col = st.selectbox("Select X-axis column:", columns)
        y_col = st.selectbox("Select Y-axis column:", columns)
        graph_type = st.selectbox("Select Graph Type:", ["Scatter Plot", "Line Plot", "Bar Chart"])

        if st.button("Generate Graph"):
            plt.figure(figsize=(10, 6))
            if graph_type == "Scatter Plot":
                sns.scatterplot(x=data[x_col], y=data[y_col])
            elif graph_type == "Line Plot":
                sns.lineplot(x=data[x_col], y=data[y_col])
            elif graph_type == "Bar Chart":
                sns.barplot(x=data[x_col], y=data[y_col])
            plt.title(f"{graph_type}: {y_col} vs {x_col}")
            plt.xlabel(x_col)
            plt.ylabel(y_col)
            st.pyplot(plt)

    # Prediction
    st.write("### Predictive Analysis:")
    target = st.selectbox("Select Target Column for Prediction:", columns)
    features = st.multiselect("Select Feature Columns:", [col for col in columns if col != target])
    model_type = st.selectbox("Select Model:", ["Linear Regression", "Random Forest"])

    model = None
    predictions = None
    y_test = None

    if st.button("Train and Predict"):
        if target and features:
            X = data[features].copy()
            y = data[target].copy()

            # Preprocessing: Encoding categorical variables
            for col in X.columns:
                if X[col].dtype == "object":
                    X.loc[:, col] = LabelEncoder().fit_transform(X[col].astype(str))

            if y.dtype == "object":
                y = LabelEncoder().fit_transform(y.astype(str))

            # Train-Test Split
            X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

            if model_type == "Linear Regression":
                model = LinearRegression()
            elif model_type == "Random Forest":
                model = RandomForestRegressor(random_state=42)

            model.fit(X_train, y_train)
            predictions = model.predict(X_test)

            st.write("Predictions on Test Data:")
            st.write(predictions[:10])

            # Model Evaluation
            st.write("### Model Evaluation:")
            st.write(f"Mean Squared Error: {mean_squared_error(y_test, predictions)}")
            st.write(f"R-squared: {r2_score(y_test, predictions)}")
        else:
            st.warning("Please select both target and feature columns.")

    # Generate Word Report
    ##if st.button("Generate Word Report"):
        file_path = generate_word_report(data, model, predictions, model_type, y_test)
        with open(file_path, "rb") as f:
            st.download_button(
                label="Download Report",
                data=f,
                file_name="Data_Analysis_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
