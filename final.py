    
import streamlit as st
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression,Ridge,Lasso
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.preprocessing import LabelEncoder
from docx import Document
import google.generativeai as genai
import re
from sklearn.tree import DecisionTreeRegressor
# Function to generate the report
from docx import Document
import matplotlib.pyplot as plt
import seaborn as sns
import io
from pymongo import MongoClient

st.set_page_config(page_title="AI-Driven Data Analysis Dashboard", layout="wide")
st.title("AI-Driven Data Analysis Dashboard")
st.sidebar.title("Navigation")

# Sidebar Navigation with Radio Button
menu = st.sidebar.radio("Go to", ["Upload Data", "Dataset Overview", "Data Visualization", "Model Training", "Generate Report", "Dashboard", "Chat interface"])

api_key = 'AIzaSyDX3HE-dhk-0xUc7amKaIz8avJ6gpUFeGo'  # Replace with your actual API key

# Configure the API with the provided key
genai.configure(api_key=api_key)

# Initialize the model
model = genai.GenerativeModel(model_name="gemini-1.5-flash")

# Function to generate content based on a prompt
def generate_content(prompt, max_words):
    response = model.generate_content([f"{prompt} (Max words: {max_words})"])
    return response.text



#css



from docx import Document
import pandas as pd
import io
import matplotlib.pyplot as plt
import seaborn as sns
import streamlit as st

def generate_word_report(data):
    doc = Document()
    doc.add_heading('Data Analysis Report', 0)

    # Dataset Overview
    doc.add_heading('Dataset Overview', level=1)
    doc.add_paragraph(f"Number of rows: {data.shape[0]}")
    doc.add_paragraph(f"Number of columns: {data.shape[1]}")
    doc.add_heading('Column Information', level=2)
    doc.add_paragraph(str(data.dtypes))

    # Dataset Preview
    doc.add_heading('Dataset Preview', level=2)
    doc.add_paragraph(f"First few rows of the dataset:\n{data.head()}")

    # Dataset Statistics in Tabular Format
    doc.add_heading('Dataset Statistics', level=1)

    # Get the description of the dataset
    desc = data.describe()

    # Add a table with the description data
    table = doc.add_table(rows=1, cols=len(desc.columns) + 1)

    # Add headers to the table
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Statistic'
    for i, col in enumerate(desc.columns):
        hdr_cells[i + 1].text = col

    # Add data to the table
    for stat in desc.index:
        row_cells = table.add_row().cells
        row_cells[0].text = stat
        for i, value in enumerate(desc.loc[stat]):
            row_cells[i + 1].text = str(value)

    # Data Transformations
    if "transformations" in st.session_state:
        doc.add_heading('Data Transformations', level=1)
        for transformation in st.session_state['transformations']:
            doc.add_paragraph(transformation)

    # Graphs and Visualizations
    if "graphs" in st.session_state and len(st.session_state['graphs']) > 0:
        doc.add_heading('Graphs and Visualizations', level=1)
        for idx, graph in enumerate(st.session_state['graphs']):
            x_col = graph["x_col"]
            y_col = graph["y_col"]
            visualization_type = graph["type"]

            doc.add_heading(f"{visualization_type}: {y_col} vs {x_col}", level=2)

            # Save the graph as an image and add it to the Word report
            fig, ax = plt.subplots(figsize=(10, 6))
            try:
                if visualization_type == "Scatter Plot":
                    sns.scatterplot(x=data[x_col], y=data[y_col], ax=ax)
                elif visualization_type == "Line Plot":
                    sns.lineplot(x=data[x_col], y=data[y_col], ax=ax)
                elif visualization_type == "Bar Chart":
                    sns.barplot(x=data[x_col], y=data[y_col], ax=ax)
                elif visualization_type == "Heatmap":
                    corr_matrix = data.corr()
                    sns.heatmap(corr_matrix, annot=True, cmap='coolwarm', ax=ax)
                elif visualization_type == "Histogram":
                    sns.histplot(data[y_col], kde=True, ax=ax)
                elif visualization_type == "Box Plot":
                    sns.boxplot(x=data[x_col], y=data[y_col], ax=ax)
                elif visualization_type == "Violin Plot":
                    sns.violinplot(x=data[x_col], y=data[y_col], ax=ax)
                elif visualization_type == "Correlation Heatmap":
                    corr_matrix = data.corr()
                    sns.heatmap(corr_matrix, annot=True, cmap='coolwarm')
                    plt.title("Correlation Heatmap")
                
                else:
                    doc.add_paragraph(f"Unsupported graph type: {visualization_type}")
                    continue

                # Set the title for the graph
                ax.set_title(f"{visualization_type}: {y_col} vs {x_col}")

                # Save the plot to a buffer
                image_stream = io.BytesIO()
                plt.savefig(image_stream, format='png')
                plt.close(fig)
                image_stream.seek(0)

                # Add the plot image to the Word document
                doc.add_picture(image_stream)
            except Exception as e:
                doc.add_paragraph(f"Error generating graph {visualization_type}: {y_col} vs {x_col}. Error: {str(e)}")
    else:
        doc.add_paragraph("No graphs have been created yet.")

    # Model Training Results
    if "model_training" in st.session_state:
        doc.add_heading('Model Training Results', level=1)
        model_info = st.session_state['model_training']
        doc.add_paragraph(f"Model Type: {model_info['model_type']}")
        doc.add_paragraph(f"Model Performance: {model_info['performance']}")
        doc.add_paragraph(f"Hyperparameters: {model_info['hyperparameters']}")

    # Predictions
    if "predictions" in st.session_state:
        doc.add_heading('Predictions', level=1)
        predictions = st.session_state['predictions']
        doc.add_paragraph(f"Predicted Values:\n{predictions}")

    # Chat Interface Logs
    if "chat_logs" in st.session_state:
        doc.add_heading('Chat Interface Logs', level=1)
        for log in st.session_state['chat_logs']:
            doc.add_paragraph(log)

    # Save the document
    file_path = "data_analysis_report.docx"
    doc.save(file_path)
    return file_path


# Handling 'Upload Data'
if menu == "Upload Data":
    st.header("Upload Your Dataset or Fetch from a Database")
    
    # Option to choose data source
    data_source = st.radio("Select Data Source", ["CSV/JSON File", "MongoDB Database"])
    
    # Handling file upload
    if data_source == "CSV/JSON File":
        uploaded_file = st.file_uploader("Upload your dataset (CSV or JSON)", type=["csv", "json"])
        if uploaded_file:
            try:
                if uploaded_file.name.endswith(".csv"):
                    data = pd.read_csv(uploaded_file)
                elif uploaded_file.name.endswith(".json"):
                    data = pd.read_json(uploaded_file)
                st.session_state['data'] = data
                st.success("File uploaded successfully!")
                st.write("### Dataset Preview:")
                st.write(data.head())
            except Exception as e:
                st.error(f"Error processing the file: {e}")
    
    # Handling MongoDB fetch
    elif data_source == "MongoDB Database":
        mongodb_uri = st.text_input("Enter MongoDB URI:")
        database_name = st.text_input("Enter Database Name:")
        collection_name = st.text_input("Enter Collection Name:")
        
        if st.button("Fetch Data from MongoDB"):
            try:
                # Connect to MongoDB
                client = MongoClient(mongodb_uri)
                database = client[database_name]
                collection = database[collection_name]
                
                # Fetch data
                documents = list(collection.find())
                if documents:
                    data = pd.DataFrame(documents)
                    st.session_state['data'] = data
                    st.success("Data fetched successfully from MongoDB!")
                    st.write("### Dataset Preview:")
                    st.write(data.head())
                else:
                    st.warning("No data found in the collection.")
            except Exception as e:
                st.error(f"Error connecting to MongoDB: {e}")
    




# Handling 'Dataset Overview

if "data" in st.session_state and menu == "Dataset Overview":
    data = st.session_state['data']
    st.header("Dataset Overview")

    st.write("### Sanity Checks")
    st.write(f"Missing Values:\n{data.isnull().sum()}")
    st.write(f"Duplicate Rows: {data.duplicated().sum()}")

    if st.button("Handle Missing Values"):
        data = data.dropna()
        st.session_state['data'] = data
        st.write("Missing values removed.")

    st.write("### Dataset Statistics")
    st.write(data.describe())


    # Handling 'Data Visualization'
if "data" in st.session_state and menu == "Data Visualization":
    data = st.session_state['data']
    st.header("Data Visualization")

    # Using Radio Buttons for selecting visualization type with a unique key
    visualization_type = st.radio(
        "Select Visualization Type", 
        [
            "Scatter Plot", "Line Plot", "Bar Chart", "Heatmap", 
            "Histogram", "Box Plot", "Violin Plot",  
            "Correlation Heatmap"
        ], 
        key="visualization_type"
    )

    columns = data.columns.tolist()

    # User selects X and Y columns for scatter, line, and bar charts
    x_col = st.selectbox("Select X-axis column:", columns, key="x_col")
    y_col = st.selectbox("Select Y-axis column:", columns, key="y_col")

    # Initialize a list in session_state to store graphs if not already initialized
    if "graphs" not in st.session_state:
        st.session_state['graphs'] = []

    if st.button("Generate Graph", key="generate_graph"):
        plt.figure(figsize=(10, 6))

        # Scatter Plot
        if visualization_type == "Scatter Plot":
            sns.scatterplot(x=data[x_col], y=data[y_col])
            plt.title(f"Scatter Plot: {y_col} vs {x_col}")
        
        # Line Plot
        elif visualization_type == "Line Plot":
            sns.lineplot(x=data[x_col], y=data[y_col])
            plt.title(f"Line Plot: {y_col} vs {x_col}")
        
        # Bar Chart
        elif visualization_type == "Bar Chart":
            sns.barplot(x=data[x_col], y=data[y_col])
            plt.title(f"Bar Chart: {y_col} vs {x_col}")
        
        elif visualization_type == "Heatmap":
            numeric_data = data.select_dtypes(include=['number'])
            if not numeric_data.empty:
                corr_matrix = numeric_data.corr()
                sns.heatmap(corr_matrix, annot=True, cmap='coolwarm')
                plt.title("Heatmap")
                plt.show()
            else:
                print("No numeric columns available for correlation calculation.")

        # Histogram
        elif visualization_type == "Histogram":
            sns.histplot(data[y_col], kde=True)
            plt.title(f"Histogram: {y_col}")
        
        # Box Plot
        elif visualization_type == "Box Plot":
            sns.boxplot(x=data[x_col], y=data[y_col])
            plt.title(f"Box Plot: {y_col} vs {x_col}")
        
        # Violin Plot
        elif visualization_type == "Violin Plot":
            sns.violinplot(x=data[x_col], y=data[y_col])
            plt.title(f"Violin Plot: {y_col} vs {x_col}")
        
        # Correlation Heatmap (only numerical columns)
        elif visualization_type == "Correlation Heatmap":
            numeric_data = data.select_dtypes(include=['number'])
            if not numeric_data.empty:
                corr_matrix = numeric_data.corr()
                sns.heatmap(corr_matrix, annot=True, cmap='coolwarm')
                plt.title("Correlation Heatmap")
                plt.show()
            else:
                print("No numeric columns available for correlation calculation.")

        plt.xlabel(x_col)
        plt.ylabel(y_col)
        st.pyplot(plt)

        # Save the current graph details to session_state
        st.session_state['graphs'].append({
            "type": visualization_type,
            "x_col": x_col,
            "y_col": y_col
        })



#word documetnx 
if "data" in st.session_state and menu == "Generate Report":
    data = st.session_state['data']
    st.header("Generate Report")
    if st.button("Download Word Report"):
        file_path = generate_word_report(data)
        with open(file_path, "rb") as f:
            st.download_button(
                label="Download Report",
                data=f,
                file_name="Data_Analysis_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


# Handling 'Model Training'
if "data" in st.session_state and menu == "Model Training":
    data = st.session_state['data']
    st.header("Model Training")
    columns = data.columns.tolist()

    # Step 1: Select Target and Features
    target = st.selectbox("Select Target Column for Prediction:", columns)
    features = st.multiselect("Select Feature Columns:", [col for col in columns if col != target])

    if st.button("Train and Evaluate Models"):
        if target and features:
            # Step 2: Data Preparation
            X = data[features].copy()
            y = data[target].copy()

            # Encode Categorical Features
            for col in X.columns:
                if X[col].dtype == "object":
                    X[col] = LabelEncoder().fit_transform(X[col].astype(str))
            if y.dtype == "object":
                y = LabelEncoder().fit_transform(y.astype(str))

            # Split the Data
            X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

            # Define Models to Evaluate
            models = {
                "Linear Regression": LinearRegression(),
                "Ridge Regression": Ridge(random_state=42),
                "Lasso Regression": Lasso(random_state=42),
                "Random Forest": RandomForestRegressor(random_state=42),
                "Decision Tree": DecisionTreeRegressor(random_state=42),
            }

            # Evaluate Models
            results = {}
            for model_name, model in models.items():
                model.fit(X_train, y_train)
                predictions = model.predict(X_test)
                mse = mean_squared_error(y_test, predictions)
                r2 = r2_score(y_test, predictions)
                results[model_name] = {"MSE": mse, "R2": r2}

            # Display Results with Comments on Performance
            #'''st.write("### Model Performance")
            #for model_name, metrics in results.items():
                #st.write(f"**{model_name}**")
                #st.write(f"Mean Squared Error: {metrics['MSE']:.2f} - This metric indicates the average squared difference between actual and predicted values. Lower MSE indicates better model performance.")
                #st.write(f"R-squared: {metrics['R2']:.2f} - This metric shows how well the model explains the variability of the target variable. Higher R² indicates better fit.")
                #st.write("---")'''

            # Find the Best Model
            best_model_name = min(results, key=lambda name: results[name]["MSE"])
            best_model = models[best_model_name]
            st.success(f"The best model is: {best_model_name} with MSE: {results[best_model_name]['MSE']:.2f}")

            # Save the Best Model in Session State
            st.session_state["trained_model"] = best_model
            st.session_state["features"] = features
            st.session_state["target"] = target
            st.success("Model training and evaluation completed!")

    # Step 3: User Input for Prediction
    st.write("### Predict Based on Feature Values")
    if "trained_model" in st.session_state and "features" in st.session_state:
        input_data = {}
        for feature in st.session_state["features"]:
            if data[feature].dtype == "object":
                unique_values = data[feature].unique()
                input_data[feature] = st.selectbox(f"Select value for {feature}:", unique_values)
            else:
                default_value = data[feature].mean()
                input_data[feature] = st.number_input(
                    f"Enter value for {feature}:",
                    value=float(default_value),
                    format="%.2f"
                )
        # Handle Prediction
        if st.button("Predict Target Value"):
            # Ensure all input data is present
            if all(input_data.values()):
                input_df = pd.DataFrame([input_data])

                # Encode Categorical Inputs
                for col in input_df.columns:
                    if data[col].dtype == "object":
                        encoder = LabelEncoder()
                        encoder.fit(data[col].astype(str))
                        input_df[col] = encoder.transform(input_df[col].astype(str))
                    else:
                        input_df[col] = input_df[col].astype(float)

                # Predict using the Trained Model
                model = st.session_state["trained_model"]
                user_prediction = model.predict(input_df)[0]
                st.success(f"Predicted {target}: {user_prediction}")
            else:
                st.error("Please fill in all feature values!")
    else:
        st.warning("Please train the models first to enable predictions.")



# Handling 'Dashboard'
if "data" in st.session_state and menu == "Dashboard":
    data = st.session_state['data']
    st.header("Dashboard")
    st.write("### Overview of All Created Graphs")

    # Check if any graphs were created in Data Visualization
    if "graphs" in st.session_state and len(st.session_state['graphs']) > 0:
        st.write("#### Individual Graphs")
        # Render each saved graph individually
        for idx, graph in enumerate(st.session_state['graphs']):
            x_col = graph["x_col"]
            y_col = graph["y_col"]
            visualization_type = graph["type"]

            plt.figure(figsize=(10, 6))
            if visualization_type == "Scatter Plot":
                sns.scatterplot(x=data[x_col], y=data[y_col])
            elif visualization_type == "Line Plot":
                sns.lineplot(x=data[x_col], y=data[y_col])
            elif visualization_type == "Bar Chart":
                sns.barplot(x=data[x_col], y=data[y_col])
            elif visualization_type == "Heatmap":
                numeric_data = data.select_dtypes(include=['number'])
                if not numeric_data.empty:
                    corr_matrix = numeric_data.corr()
                    sns.heatmap(corr_matrix, annot=True, cmap='coolwarm')
                    plt.title("Correlation Heatmap")
                    plt.show()
                else:
                    print("No numeric columns available for correlation calculation.")

            elif visualization_type == "Histogram":
                sns.histplot(data[y_col], kde=True)
            elif visualization_type == "Box Plot":
                sns.boxplot(x=data[x_col], y=data[y_col])
            elif visualization_type == "Violin Plot":
                sns.violinplot(x=data[x_col], y=data[y_col])
           
            elif visualization_type == "Correlation Heatmap":
                numeric_data = data.select_dtypes(include=['number'])
                if not numeric_data.empty:
                    corr_matrix = numeric_data.corr()
                    sns.heatmap(corr_matrix, annot=True, cmap='coolwarm')
                    plt.title("Correlation Heatmap")
                    plt.show()
                else:
                    print("No numeric columns available for correlation calculation.")

            plt.title(f"{visualization_type}: {y_col} vs {x_col}")
            plt.xlabel(x_col)
            plt.ylabel(y_col)
            st.pyplot(plt)

        # Render all graphs together in a single plot (Dashboard view)
        st.write("#### Combined Dashboard of All Graphs")
        num_graphs = len(st.session_state['graphs'])
        cols = 2  # Number of columns for the combined dashboard
        rows = -(-num_graphs // cols)  # Calculate rows needed

        fig, axes = plt.subplots(rows, cols, figsize=(10, 6))
        axes = axes.flatten()  # Flatten for easy iteration

        for idx, graph in enumerate(st.session_state['graphs']):
            x_col = graph["x_col"]
            y_col = graph["y_col"]
            visualization_type = graph["type"]

            if visualization_type == "Scatter Plot":
                sns.scatterplot(ax=axes[idx], x=data[x_col], y=data[y_col])
            elif visualization_type == "Line Plot":
                sns.lineplot(ax=axes[idx], x=data[x_col], y=data[y_col])
            elif visualization_type == "Bar Chart":
                sns.barplot(ax=axes[idx], x=data[x_col], y=data[y_col])
            elif visualization_type == "Heatmap":
                numeric_data = data.select_dtypes(include=['number'])
                if not numeric_data.empty:
                    corr_matrix = numeric_data.corr()
                    sns.heatmap(corr_matrix, annot=True, cmap='coolwarm', ax=axes[idx])
                    plt.title("Correlation Heatmap")
                else:
                    print("No numeric columns available for correlation calculation.")
            elif visualization_type == "Histogram":
                sns.histplot(data[y_col], kde=True, ax=axes[idx])
            elif visualization_type == "Box Plot":
                sns.boxplot(x=data[x_col], y=data[y_col], ax=axes[idx])
            elif visualization_type == "Violin Plot":
                sns.violinplot(x=data[x_col], y=data[y_col], ax=axes[idx])
            
            elif visualization_type == "Correlation Heatmap":
                numeric_data = data.select_dtypes(include=['number'])
                if not numeric_data.empty:
                    corr_matrix = numeric_data.corr()
                    sns.heatmap(corr_matrix, annot=True, cmap='coolwarm', ax=axes[idx])
                    axes[idx].set_title("Correlation Heatmap")
                else:
                    print("No numeric columns available for correlation calculation.")

            axes[idx].set_title(f"{visualization_type}: {y_col} vs {x_col}")
            axes[idx].set_xlabel(x_col)
            axes[idx].set_ylabel(y_col)

        # Remove unused subplots
        for ax in axes[num_graphs:]:
            ax.remove()

        plt.tight_layout()
        st.pyplot(fig)
    else:
        st.write("No graphs have been created yet in Data Visualization.")


# Handling 'Chat Interface'
# Handling 'Chat Interface'
if menu == "Chat interface":
    # Initialize Gemini AI

    # Streamlit app layout
    st.title("Data Chat Interface with Gemini AI")

    # Store uploaded data and chat history in session state
    if 'data' not in st.session_state:
        st.session_state['data'] = None
    if 'history' not in st.session_state:
        st.session_state['history'] = []

    # Helper function to display chat messages
    def display_chat():
        for message in st.session_state.history:
            if message.startswith("User:"):
                st.markdown(f"**User**: {message[6:]}")
            elif message.startswith("AI:"):
                st.markdown(f"**AI**: {message[4:]}", unsafe_allow_html=True)

    # File uploader for the dataset with a unique key

    st.write("### Dataset Preview:")
    if st.session_state['data'] is not None:
        st.write(st.session_state['data'].head())

    # Chat interface
    if st.session_state['data'] is not None:
        st.header("Chat Interface with Your Data")
        user_question = st.text_input("Ask a question about your data:")

        if user_question:
            # Add user question to history
            st.session_state.history.append(f"User: {user_question}")

            # Prepare the raw data and question for Gemini AI
            dataset_raw = st.session_state['data'].to_json(orient="records")  # Convert the dataset to a JSON format
            prompt = f"""
            Here is the raw dataset in JSON format:
            {dataset_raw}

            Answer the following question based on this data:
            {user_question}
            """

            try:
                # Use generate_content() as per your request
                response = model.generate_content([f"{prompt} (Max words: 200)"])

                # Access the output using the 'text' attribute
                ai_response = response.text.strip() if response else "Sorry, I couldn't process your query."

            except Exception as e:
                ai_response = f"Error occurred: {e}"

            # Add AI response to history
            st.session_state.history.append(f"AI: {ai_response}")

            # Display the chat
            display_chat()

        # Add a download button to save chat history
        if st.session_state.history:
            chat_history_text = "\n".join(st.session_state.history)
            st.download_button(
                label="Download Chat History",
                data=chat_history_text,
                file_name="chat_history.txt",
                mime="text/plain"
            )
